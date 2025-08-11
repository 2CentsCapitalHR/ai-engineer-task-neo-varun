from io import BytesIO
from typing import Union, Any

from docx import Document


def read_docx_text(uploaded_file: Union[BytesIO, Any]) -> str:
    # Streamlit UploadedFile supports getvalue(); BytesIO works with bytes
    data = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
    doc = Document(BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)
