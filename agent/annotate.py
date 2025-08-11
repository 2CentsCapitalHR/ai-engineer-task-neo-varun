"""
Document annotation functionality for adding review comments to DOCX files.
"""

from io import BytesIO
from typing import List, Dict

from docx import Document


def annotate_docx(filename: str, uploaded_file_obj, comments: List[Dict], rag) -> bytes:
    """
    Add review comments to a DOCX document.
    
    Args:
        filename: Name of the file being annotated
        uploaded_file_obj: Streamlit uploaded file object
        comments: List of comment dictionaries with issue details
        rag: RAG system (unused but kept for compatibility)
        
    Returns:
        Bytes of the annotated DOCX document
    """
    # Load the original document
    data = uploaded_file_obj.getvalue()
    doc = Document(BytesIO(data))

    # Ensure document has at least one paragraph
    if not doc.paragraphs:
        doc.add_paragraph("")

    # Add review comments section
    doc.add_paragraph("--- ADGM Compliance Review Comments ---")
    
    for comment in comments:
        severity = comment.get('severity', 'Unknown')
        section = comment.get('section', 'General')
        issue = comment.get('issue', 'No issue specified')
        suggestion = comment.get('suggestion', 'No suggestion provided')
        citation = comment.get('citation', '')
        
        comment_text = (
            f"[{severity}] {section}: {issue}\n"
            f"Suggestion: {suggestion}\n"
            f"{citation}"
        )
        doc.add_paragraph(comment_text)

    # Save to bytes and return
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
